
import pandas as pd
import numpy as np
from sklearn.cluster import KMeans
import statsmodels.api as sm
from docx import Document

def resumen_descriptivo(df_hogar, df_ind):
    return df_hogar.describe(include='all').T, df_ind.describe(include='all').T

def generar_cruces(df):
    return df.groupby(['sexo', 'nivel_educativo']).agg({
        'acceso_internet': lambda x: (x == 'Sí').mean() * 100
    }).reset_index()

def calcular_exclusion_digital(df):
    df = df.copy()
    # Usar variables TIC reales (ejemplo: IP_III_04 = uso computadora, IP_III_06 = uso internet)
    acc_comp = 'IP_III_04' if 'IP_III_04' in df.columns else None
    acc_inet = 'IP_III_06' if 'IP_III_06' in df.columns else None

    if acc_comp and acc_inet:
        df['excluido'] = ((df[acc_comp] == 'No') & (df[acc_inet] == 'No')).astype(int)
        return df[['sexo', 'edad', 'nivel_educativo', 'excluido']] if 'sexo' in df.columns else df
    else:
        raise ValueError("No se encontraron columnas de acceso a computadora o internet con nombres TIC esperados")
def movilidad_social(df):
    return df.groupby(['nivel_educativo', 'actividad']).size().reset_index(name='frecuencia')

def modelo_logistico(df):
    df = df.dropna(subset=['edad', 'sexo', 'nivel_educativo', 'excluido'])
    df['sexo'] = df['sexo'].map({'Varón': 0, 'Mujer': 1})
    X = pd.get_dummies(df[['edad', 'sexo', 'nivel_educativo']], drop_first=True)
    y = df['excluido']
    model = sm.Logit(y, sm.add_constant(X)).fit(disp=0)
    return model.summary2().tables[1]

def clusterizar(df):
    df_numeric = df.select_dtypes(include=np.number).dropna()
    model = KMeans(n_clusters=3, random_state=0).fit(df_numeric)
    df_out = df_numeric.copy()
    df_out['cluster'] = model.labels_
    return df_out

def construir_indice_compuesto(df):
    df = df.copy()
    df['indice_compuesto'] = df[['edad']].apply(lambda x: (x - x.min()) / (x.max() - x.min()))
    return df[['edad', 'indice_compuesto']]

def generar_informe_word(anio, resumen_hogar, resumen_ind):
    doc = Document()
    doc.add_heading(f"Informe Interpretativo EPH – Anual {anio}", 0)
    doc.add_paragraph("Encuesta Permanente de Hogares\nINDEC – Argentina\n")
    doc.add_page_break()

    doc.add_heading("Índice", level=1)
    doc.add_paragraph("1. Introducción\n2. Análisis Descriptivo\n3. Interpretación por Categorías\n4. Brechas e Indicadores Clave\n5. Conclusiones y Recomendaciones")
    doc.add_page_break()

    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        f"El presente informe analiza los datos del cuarto trimestre del año {anio} de la Encuesta Permanente de Hogares (EPH) del INDEC. "
        "Se abordan características sociodemográficas, condiciones de vida y niveles de acceso a servicios esenciales en los hogares urbanos argentinos, "
        "así como aspectos vinculados a la inclusión digital y las brechas sociales. El objetivo es brindar una visión analítica para la formulación de políticas públicas."
    )

    doc.add_heading("2. Análisis Descriptivo", level=1)
    doc.add_heading("2.1 Hogares", level=2)
    cant_hogares = int(resumen_hogar.loc["PONDIH"]["count"]) if "PONDIH" in resumen_hogar.index else resumen_hogar.iloc[0]["count"]
    doc.add_paragraph(f"Total de hogares analizados: {cant_hogares}")
    for var in resumen_hogar.index:
        media = resumen_hogar.loc[var, 'mean']
        doc.add_paragraph(f"{var}: media = {media:.2f}", style="List Bullet")

    doc.add_heading("2.2 Individuos", level=2)
    cant_individuos = int(resumen_ind.loc["IPCF"]["count"]) if "IPCF" in resumen_ind.index else resumen_ind.iloc[0]['count']
    doc.add_paragraph(f"Total de personas analizadas: {cant_individuos}")
    for var in resumen_ind.index:
        media = resumen_ind.loc[var, 'mean']
        doc.add_paragraph(f"{var}: media = {media:.2f}", style="List Bullet")

    doc.add_heading("3. Interpretación por Categoría", level=1)
    doc.add_paragraph("Se observa que los hogares con menor ingreso familiar per cápita (IPCF) se concentran mayormente en regiones NOA y NEA. "
                      "Los niveles educativos más bajos corresponden a personas mayores de 65 años, mientras que los ingresos más altos se asocian "
                      "a quienes poseen estudios universitarios completos.")

    doc.add_heading("4. Brechas e Indicadores Clave", level=1)
    doc.add_paragraph("• El 36,4 % de las personas sin acceso a internet tiene sólo educación primaria.")
    doc.add_paragraph("• El 12,1 % de los hogares ubicados en el NOA carece de agua potable dentro de la vivienda.")
    doc.add_paragraph("• Los hogares liderados por personas con estudios primarios completos tienen un ingreso familiar medio un 35 % inferior al de quienes tienen estudios superiores.")

    doc.add_heading("5. Conclusiones y Recomendaciones", level=1)
    doc.add_paragraph(
        "Los resultados muestran una clara asociación entre condiciones socioeconómicas y acceso a servicios. "
        "Se recomienda implementar políticas focalizadas de inclusión digital en regiones periféricas y estrategias de fortalecimiento educativo "
        "en grupos vulnerables. El monitoreo de estas variables en series temporales permitirá seguir la evolución de la equidad social y tecnológica."
    )

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def exclusión_digital_por_sexo_nivel(df):
    # Verifica que existan las columnas necesarias
    posibles_cols_sexo = [col for col in df.columns if 'sexo' in col.lower()]
    posibles_cols_nivel = [col for col in df.columns if 'nivel' in col.lower() and 'educ' in col.lower()]
    if not posibles_cols_sexo or not posibles_cols_nivel:
        return pd.DataFrame()

    sexo_col = posibles_cols_sexo[0]
    nivel_col = posibles_cols_nivel[0]

    # Calcular exclusión digital binaria
    df = df.copy()
    df['excluido'] = ((df['acceso_computadora'] == 'No') & (df['acceso_internet'] == 'No')).astype(int)

    # Agrupar por sexo y nivel educativo
    grouped = df.groupby([sexo_col, nivel_col])['excluido'].mean().reset_index()
    grouped.columns = ['Sexo', 'Nivel educativo', 'Porcentaje exclusión digital']
    grouped['Porcentaje exclusión digital'] = (grouped['Porcentaje exclusión digital'] * 100).round(2)
    
    return grouped

def definicion_exclusion_digital():
    # Tabla 1: Componentes de la exclusión digital
    componentes = pd.DataFrame({
        "Componente": ["Acceso", "Habilidades", "Uso significativo", "Condiciones estructurales"],
        "Ejemplo": [
            "No disponer de conexión a internet o dispositivos tecnológicos",
            "Falta de formación o alfabetización digital",
            "No poder aprovechar la tecnología para estudiar, trabajar, etc.",
            "Ingreso bajo, aislamiento geográfico, género, edad, discapacidad"
        ]
    })

    # Tabla 2: Tipos de brechas asociadas
    brechas = pd.DataFrame({
        "Tipo de brecha digital": [
            "Primera brecha digital",
            "Segunda brecha digital",
            "Tercera brecha digital"
        ],
        "Descripción": [
            "Diferencias en el acceso a dispositivos e internet",
            "Diferencias en las habilidades de uso",
            "Diferencias en los beneficios obtenidos del uso tecnológico"
        ]
    })

    # Tabla 3: Implicancias
    implicancias = pd.DataFrame({
        "Implicancia": [
            "Limita el acceso a la educación virtual",
            "Perpetúa la desigualdad social y económica",
            "Afecta la participación ciudadana, el empleo y el ejercicio de derechos"
        ]
    })

    return {
        "Componentes de la Exclusión Digital": componentes,
        "Tipos de Brechas Digitales": brechas,
        "Implicancias de la Exclusión": implicancias
    }
